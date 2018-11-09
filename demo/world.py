import docx
import os
import re


# world操作
# 传入三个参数, 旧字符串, 新字符串, 文件对象
def replace_text(old_text, new_text, file):
    # 遍历文件对象
    for f in file.paragraphs:
        # 如果 旧字符串 在 某个段落 中
        if old_text in f.text:
            print("替换前:", f.text)
            # 将段落存入 inline
            inline = f.runs
            # 遍历 段落 生成 i
            for i in inline:
                # 如果 旧字符串 在 i 中
                if old_text in i.text:
                    # 替换 i.text 内文本资源
                    text = i.text.replace(old_text, new_text)
                    i.text = text
            print("替换后===>", f.text)


def main():
    # 获取当前目录下所有的文件名列表
    old_file_names = os.listdir()

    # 获取所有docx文件名列表
    docx_file_names = []
    for old_file_name in old_file_names:
        if re.match(r'^[^~].*\.docx', old_file_name):
            print(old_file_name)
            docx_file_names.append(old_file_name)

    for docx_file_name in docx_file_names:
        try:
            # 获取文件对象
            file = docx.Document(docx_file_name)
            # 三个参数: 旧的字符串, 新的字符串, 文件对象
            print("开始替换:", docx_file_name)
            replace_text('海南大学', 'Hainan University', file)
            file.save('new_' + docx_file_name)
            print(docx_file_name, "替换成功")
        except:
            print(docx_file_name, "替换失败")
            pass


if __name__ == '__main__':
    docx_file_name = 'D:/tt/模板.docx'
    # 获取文件对象
    file = docx.Document(docx_file_name)
    # 三个参数: 旧的字符串, 新的字符串, 文件对象
    print("开始替换:", docx_file_name)
    replace_text('天气', '晴天 2-10 北风', file)
    replace_text('本月余额', '本月余额22天', file)
    file.save(docx_file_name)
    print(docx_file_name, "替换成功")
    # main()
