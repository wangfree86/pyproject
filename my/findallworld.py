import docx

# 读取word内容
#　这里是以段落为单位的，下面用一个for 遍历所有段落
url='D:\BaiduNetdiskDownload\\tt.docx'
doc = docx.Document(url)
doc.paragraphs
parag_num = 0
for para in doc.paragraphs :
    t=para.text
    t=t.split('(')
    if len(t)==2:
        print(t[0])
        print(t[1][:-1])
        parag_num += 1
print ('This document has ', parag_num, ' paragraphs')


